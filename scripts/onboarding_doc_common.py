#!/usr/bin/env python3
"""Generate AXIOM-styled employee onboarding Word document (beginner-friendly)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "cursor" / "onboarding" / "AXIOM_EMPLOYEE_ONBOARDING.docx"

# AXIOM tokens
BLACK = RGBColor(0x08, 0x08, 0x08)
INK_PRIMARY = RGBColor(0xF0, 0xF0, 0xF0)
INK_SECONDARY = RGBColor(0xC4, 0xC4, 0xC4)
INK_MUTED = RGBColor(0x9A, 0x9A, 0x9A)
INK_FAINT = RGBColor(0x6E, 0x6E, 0x6E)
BODY_DARK = RGBColor(0x22, 0x22, 0x22)
BODY_MID = RGBColor(0x55, 0x55, 0x55)
COMMAND_LIVE = RGBColor(0x4A, 0x9E, 0xFF)
COMMAND_STABLE = RGBColor(0x3D, 0xD6, 0x8C)
COMMAND_WATCH = RGBColor(0xE8, 0xA8, 0x38)

FONT_SANS = "Calibri"
FONT_DISPLAY = "Arial"
FONT_MONO = "Consolas"


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text: str, url: str, hex_color: str = "4A9EFF") -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), hex_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color_el)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _run_font(run, name: str, size_pt: float, color: RGBColor, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_spacer(doc: Document, pts: float = 12) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_section_divider(doc: Document, title: str, subtitle: str = "") -> None:
    """Thin gray bar between subsections within a day."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "CCCCCC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    if title:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title.upper())
        _run_font(r, FONT_MONO, 8, BODY_MID, bold=True)
    add_spacer(doc, 10)


def add_day_opener(
    doc: Document,
    day_label: str,
    title: str,
    meta: str,
    note: str = "",
) -> None:
    """Full page-break section header: DAY N — everything until the next day belongs here."""
    add_page_break(doc)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "080808")
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "16")
        el.set(qn("w:color"), "1A1A1A")
        borders.append(el)
    tc_pr.append(borders)

    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(20)
    p0.paragraph_format.space_after = Pt(8)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run("AXIOM  |  EMPLOYEE ONBOARDING")
    _run_font(r0, FONT_MONO, 9, INK_MUTED)

    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run(day_label.upper())
    _run_font(r1, FONT_DISPLAY, 32, INK_PRIMARY, bold=True)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(title)
    _run_font(r2, FONT_SANS, 15, INK_SECONDARY, bold=True)

    p3 = cell.paragraphs[0] if False else cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(18)
    r3 = p3.add_run(meta)
    _run_font(r3, FONT_MONO, 10, INK_MUTED)

    add_spacer(doc, 14)
    add_body(
        doc,
        "This section is dedicated to this day only. Complete every step below before you open the next day.",
        space_after=10,
    )
    if note:
        add_callout(doc, "Today", note, "tip")


def add_appendix_opener(doc: Document, label: str, title: str, subtitle: str = "") -> None:
    add_page_break(doc)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "111111")
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(16)
    p0.paragraph_format.space_after = Pt(6)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(label.upper())
    _run_font(r0, FONT_MONO, 11, COMMAND_LIVE, bold=True)
    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(12)
    r1 = p1.add_run(title)
    _run_font(r1, FONT_DISPLAY, 18, INK_PRIMARY, bold=True)
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(14)
        r2 = p2.add_run(subtitle)
        _run_font(r2, FONT_SANS, 10.5, INK_SECONDARY)
    add_spacer(doc, 12)


def add_section_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    _run_font(run, FONT_MONO, 9, INK_MUTED)
    run.font.all_caps = True


def add_part_heading(
    doc: Document,
    part_num: int,
    title: str,
    subtitle: str = "",
    *,
    page_break: bool = True,
) -> None:
    if page_break:
        add_page_break(doc)
    add_spacer(doc, 8)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"PART {part_num}")
    _run_font(r, FONT_MONO, 10, COMMAND_LIVE, bold=True)
    add_spacer(doc, 4)
    hp = doc.add_paragraph()
    hp.paragraph_format.space_after = Pt(6)
    hr = hp.add_run(title)
    _run_font(hr, FONT_DISPLAY, 20, BLACK, bold=True)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(14)
        sr = sp.add_run(subtitle)
        _run_font(sr, FONT_SANS, 11, BODY_MID)


def add_display_heading(doc: Document, text: str, level: int = 2) -> None:
    sizes = {2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    _run_font(run, FONT_DISPLAY, sizes.get(level, 12), BLACK, bold=True)


def add_body(doc: Document, text: str, space_after: float = 10) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    _run_font(run, FONT_SANS, 11, BODY_DARK)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    _run_font(run, FONT_SANS, 11, BODY_DARK)


def add_numbered_step(
    doc: Document,
    number: int,
    title: str,
    instructions: list[str],
    command: str | None = None,
    success: str | None = None,
) -> None:
    """One clear step: title, plain instructions, optional single command, optional success check."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"Step {number}. {title}")
    _run_font(r, FONT_SANS, 11.5, BLACK, bold=True)

    for line in instructions:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(4)
        bp.paragraph_format.left_indent = Inches(0.25)
        bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        bp.paragraph_format.line_spacing = 1.25
        br = bp.add_run(line)
        _run_font(br, FONT_SANS, 11, BODY_DARK)

    if command:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.left_indent = Inches(0.25)
        cr = cp.add_run("Copy this line, paste into PowerShell, press Enter:")
        _run_font(cr, FONT_SANS, 10, INK_MUTED, italic=True)
        add_code_block(doc, command, indent=True)

    if success:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        sp.paragraph_format.space_after = Pt(10)
        sp.paragraph_format.left_indent = Inches(0.25)
        sr1 = sp.add_run("You should see: ")
        _run_font(sr1, FONT_SANS, 10.5, COMMAND_STABLE, bold=True)
        sr2 = sp.add_run(success)
        _run_font(sr2, FONT_SANS, 10.5, BODY_DARK)


def add_callout(doc: Document, label: str, text: str, kind: str = "tip") -> None:
    colors = {"tip": "E8F4FF", "watch": "FFF8E8", "stable": "E8F8EF"}
    fill = colors.get(kind, colors["tip"])
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "CCCCCC")
        borders.append(el)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(f"{label.upper()}  ")
    color = COMMAND_WATCH if kind == "watch" else COMMAND_LIVE if kind == "tip" else COMMAND_STABLE
    _run_font(r1, FONT_MONO, 8.5, color, bold=True)
    r2 = p.add_run(text)
    _run_font(r2, FONT_SANS, 10.5, BODY_DARK)
    add_spacer(doc, 8)


def add_code_block(doc: Document, code: str, indent: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "0A0A0A")
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "1A1A1A")
        borders.append(el)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code.strip())
    _run_font(run, FONT_MONO, 10, INK_PRIMARY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    if indent:
        spacer.paragraph_format.left_indent = Inches(0.25)


def add_link_line(doc: Document, label: str, url: str, when: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(f"• {label}")
    _run_font(run, FONT_SANS, 11, BODY_DARK, bold=True)
    if when:
        run2 = p.add_run(f"  ({when})  ")
        _run_font(run2, FONT_SANS, 10, INK_MUTED)
    add_hyperlink(p, url, url)


def add_checklist_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(f"[ ]  {text}")
    _run_font(r, FONT_SANS, 11, BODY_DARK)


def add_toc_day(doc: Document, day: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(f"{day}  ")
    _run_font(r1, FONT_MONO, 10, COMMAND_LIVE, bold=True)
    r2 = p.add_run(title)
    _run_font(r2, FONT_SANS, 11, BODY_DARK)


def _day_git_push_steps(doc: Document, commit_msg: str, start: int = 1) -> int:
    steps = [
        ("Get latest from GitHub", ["Always pull before you push."], "git pull", "Already up to date or success"),
        ("Stage your files", ["This tells Git which files changed."], "git add .", "No output is normal"),
        ("Save with a message", ["Describe what you did today."], f'git commit -m "{commit_msg}"', "1 file changed or similar"),
        ("Upload to GitHub", ["Send your work to the team repo."], "git push", "Completes without red errors"),
    ]
    n = start
    for title, instr, cmd, ok in steps:
        add_numbered_step(doc, n, title, instr, command=cmd, success=ok)
        n += 1
    return n

