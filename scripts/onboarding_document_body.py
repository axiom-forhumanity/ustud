"""Day-segmented content for AXIOM employee onboarding Word document."""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.shared import Pt

from onboarding_doc_common import (
    INK_MUTED,
    INK_PRIMARY,
    INK_SECONDARY,
    BODY_DARK,
    BLACK,
    FONT_DISPLAY,
    FONT_MONO,
    FONT_SANS,
    WD_ALIGN_PARAGRAPH,
    add_appendix_opener,
    add_body,
    add_bullet,
    add_callout,
    add_checklist_item,
    add_code_block,
    add_day_opener,
    add_display_heading,
    add_hyperlink,
    add_link_line,
    add_numbered_step,
    add_page_break,
    add_section_divider,
    add_spacer,
    add_toc_day,
    _day_git_push_steps,
    _run_font,
    _set_cell_shading,
)


def populate_intro(doc: Document) -> None:
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    _set_cell_shading(header_cell, "080808")
    hp = header_cell.paragraphs[0]
    hr1 = hp.add_run("AXIOM  |  EMPLOYEE ONBOARDING\n")
    _run_font(hr1, FONT_MONO, 9, INK_MUTED)
    hr2 = hp.add_run("Two-Week Program\n")
    _run_font(hr2, FONT_DISPLAY, 22, INK_PRIMARY, bold=True)
    hr3 = hp.add_run("One section per day. Start at Day 0.")
    _run_font(hr3, FONT_SANS, 11, INK_SECONDARY)

    add_spacer(doc, 16)
    add_body(
        doc,
        "Each day begins on its own page with a DAY header. Everything until the next DAY header belongs to that day only. "
        "Follow days in order: Day 0, Day 1, Day 2, and so on.",
        space_after=14,
    )
    add_callout(
        doc,
        "Schedule",
        "Day 0: 2 to 3 hours (laptop setup). Days 1 to 10: about 1 hour each weekday across two weeks. "
        "Session recordings required Days 1 to 5 (Week 1) unless management states otherwise.",
        "tip",
    )

    add_display_heading(doc, "How to use PowerShell", 3)
    add_numbered_step(
        doc,
        1,
        "Open PowerShell",
        ["Start menu, type PowerShell, open Windows PowerShell."],
        success="Window shows PS C:\\Users\\YourName>",
    )
    add_numbered_step(
        doc,
        2,
        "Copy and paste commands",
        [
            "Copy text from the dark boxes in each day section.",
            "Right-click in PowerShell to paste (or Ctrl+V). Press Enter.",
            "Run one command at a time. Wait for each to finish.",
        ],
    )

    add_display_heading(doc, "Document map", 3)
    for day, title in [
        ("Day 0", "Laptop setup and accounts"),
        ("Day 1", "Week 1 Monday - GitHub, clone, first push, OBS setup"),
        ("Day 2", "Week 1 Tuesday - Automation and setup script"),
        ("Day 3", "Week 1 Wednesday - AI and LLMs"),
        ("Day 4", "Week 1 Thursday - Agents and Cursor"),
        ("Day 5", "Week 1 Friday - Run UStud locally"),
        ("Day 6", "Week 2 Monday - Codebase orientation"),
        ("Day 7", "Week 2 Tuesday - First development task"),
        ("Day 8", "Week 2 Wednesday - AI integration and tutor test"),
        ("Day 9", "Week 2 Thursday - Second development task"),
        ("Day 10", "Week 2 Friday - Third deliverable and summary"),
        ("Appendix A", "Cursor session commands"),
        ("Appendix B", "Free learning resources"),
        ("Appendix C", "Support and troubleshooting"),
        ("Appendix D", "Onboarding completion criteria"),
    ]:
        add_toc_day(doc, day, title)


def populate_day_0(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 0",
        "Laptop Setup and Accounts",
        "Before Week 1  |  Allow 2 to 3 hours  |  Windows 10 or 11",
        "Install software and create accounts. Complete before Day 1.",
    )

    installs = [
        ("Git", "Saves and sends work to GitHub.", "https://git-scm.com/download/win", "git --version"),
        ("Python 3.10+", "Runs the UStud backend. Check Add to PATH on install.", "https://www.python.org/downloads/", "python --version"),
        ("Node.js LTS", "Runs the UStud interface.", "https://nodejs.org/", "node --version"),
        ("Cursor", "Code editor with AI assistant.", "https://cursor.com/download", None),
        ("Ollama", "Offline AI tutor.", "https://ollama.com/download", "ollama --version"),
        ("OBS Studio", "Session recordings during Week 1.", "https://obsproject.com/", None),
    ]
    step = 1
    for name, why, url, verify in installs:
        add_section_divider(doc, f"Install {name}")
        add_body(doc, why)
        lp = doc.add_paragraph()
        lr = lp.add_run("Download: ")
        _run_font(lr, FONT_SANS, 11, BODY_DARK, bold=True)
        add_hyperlink(lp, url, url)
        add_numbered_step(doc, step, f"Download and install {name}", ["Run the installer. Default options are fine."], success=f"{name} installed")
        step += 1
        if verify:
            add_numbered_step(doc, step, "Verify install", ["In PowerShell:"], command=verify, success="Version number displayed")
            step += 1

    add_section_divider(doc, "Accounts")
    add_numbered_step(doc, step, "Create GitHub account", ["Sign up and verify email. Send username to management."])
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://github.com/signup", "https://github.com/signup")
    step += 1
    add_numbered_step(doc, step, "Accept repo invite", ["Accept axiom-forhumanity/ustud invite from email."])
    lp2 = doc.add_paragraph()
    add_hyperlink(lp2, "https://github.com/axiom-forhumanity/ustud", "https://github.com/axiom-forhumanity/ustud")
    step += 1
    add_numbered_step(doc, step, "Create NVIDIA Developer account", ["For free courses on Day 3 and Day 8."])
    lp3 = doc.add_paragraph()
    add_hyperlink(lp3, "https://developer.nvidia.com/login", "https://developer.nvidia.com/login")


def populate_day_1(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 1",
        "GitHub and First Push",
        "Week 1  |  Monday  |  ~1 hour  |  Session recording required",
        "Learn GitHub basics, clone the repo, push your first commit, and set up OBS.",
    )

    add_section_divider(doc, "GitHub concepts")
    add_body(doc, "The project lives on GitHub. You download a copy (clone), work locally, then upload changes (push).")
    for word, meaning in [
        ("Clone", "Download the project once."),
        ("Pull", "Get updates before you work each day."),
        ("Commit", "Save a snapshot with a note."),
        ("Push", "Send commits to GitHub."),
    ]:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{word}: ")
        _run_font(r1, FONT_SANS, 11, BLACK, bold=True)
        r2 = p.add_run(meaning)
        _run_font(r2, FONT_SANS, 11, BODY_DARK)

    add_display_heading(doc, "One-time Git identity", 3)
    add_code_block(doc, 'git config --global user.name "Your Full Name"')
    add_code_block(doc, 'git config --global user.email "your.email@example.com"')

    add_section_divider(doc, "Clone and open project")
    n = 1
    add_numbered_step(doc, n, "Open PowerShell", []); n += 1
    add_numbered_step(doc, n, "Go to Documents", [], "cd $HOME\\Documents", "Prompt shows Documents"); n += 1
    add_numbered_step(doc, n, "Clone the repo", ["Wait 1 to 2 minutes."], "git clone https://github.com/axiom-forhumanity/ustud.git", "Cloning into ustud..."); n += 1
    add_numbered_step(doc, n, "Enter project folder", [], "cd ustud", "Prompt shows ustud"); n += 1
    add_numbered_step(doc, n, "Open in Cursor", ["File, Open Folder, select Documents\\ustud"], success="Files visible on the left"); n += 1
    add_numbered_step(doc, n, "Update WEEK1_LOG.md", ["Open progress/WEEK1_LOG.md. Fill Monday section. Save (Ctrl+S)."]); n += 1
    _day_git_push_steps(doc, "Day1: GitHub setup and first push", start=n)

    add_section_divider(doc, "OBS setup (for daily recordings this week)")
    obs_n = 1
    add_numbered_step(doc, obs_n, "Open OBS Studio", []); obs_n += 1
    add_numbered_step(doc, obs_n, "Add Display Capture and webcam", ["Sources +, Display Capture, then Video Capture Device. Resize webcam to corner."]); obs_n += 1
    add_numbered_step(doc, obs_n, "Test record 30 seconds", ["Confirm screen and face appear in playback."])

    add_page_break(doc)
    add_display_heading(doc, "Recording upload (Days 1 to 5)", 3)
    add_numbered_step(doc, 1, "Name file Week1_DayN_YYYY-MM-DD_YourName.mp4", ["Day 1 = Monday through Day 5 = Friday"])
    add_numbered_step(doc, 2, "Upload to management folder same day", [])
    add_numbered_step(doc, 3, "Paste link in WEEK1_LOG.md under Trust recording URL", [])
    _day_git_push_steps(doc, "DayN: trust recording link added", start=4)

    add_section_divider(doc, "Learning link for today")
    add_link_line(doc, "GitHub Hello World", "https://docs.github.com/en/get-started/start-your-journey/hello-world", "Optional ~30 min")


def populate_day_2(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 2",
        "Automation Basics",
        "Week 1  |  Tuesday  |  ~1 hour  |  Session recording required",
        "Learn what automation means and run the project setup script.",
    )
    add_body(doc, "Automation = the computer runs steps for you. setup.ps1 installs packages without manual clicks.")
    add_numbered_step(doc, 1, "Start OBS recording", [])
    add_numbered_step(doc, 2, "Open PowerShell in ustud folder", [], "cd $HOME\\Documents\\ustud", "In ustud folder")
    add_numbered_step(doc, 3, "Pull latest", [], "git pull", "Success")
    add_numbered_step(doc, 4, "Run setup script", ["May take several minutes."], ".\\setup.ps1", "Success messages")
    add_numbered_step(doc, 5, "Update WEEK1_LOG.md Tuesday section", [])
    _day_git_push_steps(doc, "Day2: automation and setup script", start=6)
    add_section_divider(doc, "Learning link for today")
    add_link_line(doc, "GitHub Actions quickstart", "https://docs.github.com/en/actions/quickstart", "Optional ~20 min")


def populate_day_3(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 3",
        "AI and LLMs",
        "Week 1  |  Wednesday  |  ~1 hour  |  Session recording required",
        "Study generative AI and connect it to UStud's offline tutor.",
    )
    add_numbered_step(doc, 1, "Start OBS recording", [])
    add_numbered_step(doc, 2, "Open NVIDIA free courses and sign in", [])
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://www.nvidia.com/en-us/training/self-paced-courses/", "https://www.nvidia.com/en-us/training/self-paced-courses/")
    add_numbered_step(doc, 3, "Enroll in Generative AI Explained", ["Filter Free Courses. Do first hour today."])
    add_numbered_step(
        doc,
        4,
        "Answer in WEEK1_LOG.md (Wednesday)",
        ["What is an LLM?", "What is offline AI?", "Why does UStud use local models?"],
    )
    _day_git_push_steps(doc, "Day3: AI and LLM notes", start=5)
    add_section_divider(doc, "Learning link for today")
    add_link_line(doc, "Generative AI Explained (NVIDIA DLI)", "https://www.nvidia.com/en-us/training/self-paced-courses/", "Required")


def populate_day_4(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 4",
        "Agents and Cursor",
        "Week 1  |  Thursday  |  ~1 hour  |  Session recording required",
        "Finish AI course content and run your first Cursor session workflow.",
    )
    add_numbered_step(doc, 1, "Start OBS recording", [])
    add_numbered_step(doc, 2, "Continue NVIDIA course or start Augment Your LLM Using RAG", [])
    add_numbered_step(doc, 3, "git pull", [], "git pull", "Success")
    add_numbered_step(doc, 4, "Paste Session Start in Cursor", ["See Appendix A for the full block."])
    add_numbered_step(doc, 5, "Ask Cursor", ["Explain the folder structure of this repo in simple terms."])
    add_numbered_step(doc, 6, "Paste Session End in Cursor", ["Review progress file updates."])
    _day_git_push_steps(doc, "Day4: agents and Cursor session", start=7)
    add_section_divider(doc, "Learning links for today")
    add_link_line(doc, "Cursor Get Started", "https://docs.cursor.com/get-started/introduction", "Optional ~20 min")


def populate_day_5(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 5",
        "Run UStud Locally",
        "Week 1  |  Friday  |  ~1 hour  |  Session recording required",
        "Start the application and deliver your first small improvement.",
    )
    add_numbered_step(doc, 1, "Start OBS recording", [])
    add_numbered_step(doc, 2, "Pull latest", [], "git pull", "Success")
    add_numbered_step(doc, 3, "Install dependencies", [], "npm install", "Completes")
    add_numbered_step(doc, 4, "Install UI dependencies", [], "cd boiler", "In boiler folder")
    add_numbered_step(doc, 5, "UI npm install", [], "npm install", "Completes")
    add_numbered_step(doc, 6, "Return to root", [], "cd ..", "In ustud")
    add_numbered_step(doc, 7, "Start app", [], "npm run dev:all", "API and UI start")
    add_numbered_step(doc, 8, "Open browser", ["Go to http://localhost:5173"], success="UStud interface loads")
    add_numbered_step(
        doc,
        9,
        "Update CURRENT_STATE.md and one small fix",
        ["Top 3 gaps, what works, what does not. Session End, then push."],
    )
    _day_git_push_steps(doc, "Day5: first local run and improvement", start=10)


def populate_day_6(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 6",
        "Codebase Orientation",
        "Week 2  |  Monday  |  ~1 hour",
        "Map the repo before making changes. Session recording not required unless management asks.",
    )
    add_numbered_step(doc, 1, "git pull", [], "git pull", "Success")
    add_numbered_step(doc, 2, "Session Start in Cursor", ["Appendix A"])
    add_body(doc, "Ask Cursor to explain each folder:")
    add_bullet(doc, "src/ustud/ - Python backend")
    add_bullet(doc, "boiler/ - React UI")
    add_bullet(doc, "modules/ - course content")
    add_bullet(doc, "progress/ - logs and status")
    add_numbered_step(doc, 3, "Expand CURRENT_STATE.md", ["Full sentences for each gap. Add recommended next task."])
    add_numbered_step(doc, 4, "Update WEEK2_LOG.md Monday section", [])
    add_numbered_step(doc, 5, "Session End and push", [])
    _day_git_push_steps(doc, "Day6: codebase orientation", start=6)


def populate_day_7(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 7",
        "First Development Task",
        "Week 2  |  Tuesday  |  ~1 hour",
        "Implement improvement #1 from CURRENT_STATE.md.",
    )
    add_numbered_step(doc, 1, "git pull and Session Start", [], "git pull", "Success")
    add_numbered_step(
        doc,
        2,
        "Implement gap #1",
        ["Ask Cursor: Help me implement the first gap in CURRENT_STATE.md. Keep the change small."],
    )
    add_numbered_step(doc, 3, "Test locally if app is running", [])
    add_numbered_step(doc, 4, "Session End, update WEEK2_LOG.md", [])
    _day_git_push_steps(doc, "Day7: first development task", start=5)


def populate_day_8(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 8",
        "AI Integration",
        "Week 2  |  Wednesday  |  ~1 hour",
        "Complete NVIDIA training and test the offline tutor.",
    )
    add_numbered_step(doc, 1, "Finish NVIDIA RAG or Generative AI course", [])
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://www.nvidia.com/en-us/training/self-paced-courses/", "https://www.nvidia.com/en-us/training/self-paced-courses/")
    add_numbered_step(doc, 2, "Run app and test tutor", ["npm run dev:all, open http://localhost:5173, ask tutor one question."])
    add_numbered_step(doc, 3, "Document results in WEEK2_LOG.md", ["Did it work? What would improve it?"])
    _day_git_push_steps(doc, "Day8: AI integration and tutor test", start=4)


def populate_day_9(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 9",
        "Second Development Task",
        "Week 2  |  Thursday  |  ~1 hour",
        "UI or module improvement (gap #2 or related polish).",
    )
    add_numbered_step(doc, 1, "git pull and Session Start", [], "git pull", "Success")
    add_numbered_step(doc, 2, "Work in boiler/ or modules/", ["Verify in browser. Session End before push."])
    add_numbered_step(doc, 3, "Update WEEK2_LOG.md", [])
    _day_git_push_steps(doc, "Day9: second development task", start=4)


def populate_day_10(doc: Document) -> None:
    add_day_opener(
        doc,
        "Day 10",
        "Third Deliverable and Summary",
        "Week 2  |  Friday  |  ~1 hour",
        "Final improvement and two-week summary for management.",
    )
    add_numbered_step(doc, 1, "git pull", [], "git pull", "Success")
    add_numbered_step(doc, 2, "Third improvement or refine prior work", ["Confirm app still runs locally."])
    add_numbered_step(
        doc,
        3,
        "Two-week summary in DAILY_LOG.md",
        ["What you built in Week 2.", "Blockers.", "Recommended Week 3 focus.", "Complete WEEK2_LOG.md Friday."],
    )
    add_numbered_step(doc, 4, "Session End and push", [])
    _day_git_push_steps(doc, "Day10: third deliverable and onboarding summary", start=5)
    add_callout(doc, "Next step", "Review Appendix D with management. Then follow cursor/EMPLOYEE_ONBOARDING.md daily.", "stable")


def populate_appendix_a(doc: Document) -> None:
    add_appendix_opener(doc, "Appendix A", "Cursor Session Commands", "Reference for every work session")
    add_display_heading(doc, "Session Start", 2)
    add_code_block(
        doc,
        """Read and follow cursor/AXIOM_OPERATING_INSTRUCTIONS.md completely.

Also read:
- docs/PROJECT_BRIEF.md
- progress/CURRENT_STATE.md
- progress/BLOCKERS.md
- the latest entry in progress/DAILY_LOG.md

Confirm you understand the progress logging requirements.

Then tell me:
1. Current project state
2. Open blockers
3. What I should focus on today

Do not start coding until you have read these files.""",
    )
    add_display_heading(doc, "Session End", 2)
    add_code_block(
        doc,
        """Session ending. Follow cursor/AXIOM_OPERATING_INSTRUCTIONS.md session END workflow.

Update:
- progress/DAILY_LOG.md (append today's entry)
- progress/CURRENT_STATE.md
- progress/BLOCKERS.md (if needed)
- docs/DECISION_LOG.md (if any decisions were made)

Then show me a summary of exactly what you logged so I can review before I commit and push to GitHub.""",
    )


def populate_appendix_b(doc: Document) -> None:
    add_appendix_opener(doc, "Appendix B", "Free Learning Resources", "Hyperlinks by day")
    add_page_break(doc)
    add_display_heading(doc, "Free Learning Resources", 2)
    links = [
        ("Team project on GitHub", "https://github.com/axiom-forhumanity/ustud", "All days"),
        ("Project brief", "https://github.com/axiom-forhumanity/ustud/blob/main/docs/PROJECT_BRIEF.md", "Day 6+"),
        ("AXIOM vision", "https://github.com/axiom-forhumanity/ustud/blob/main/docs/AXIOM_VISION.md", "Day 6+"),
        ("GitHub Hello World", "https://docs.github.com/en/get-started/start-your-journey/hello-world", "Day 1"),
        ("GitHub Actions quickstart", "https://docs.github.com/en/actions/quickstart", "Day 2"),
        ("NVIDIA DLI self-paced (filter Free Courses)", "https://www.nvidia.com/en-us/training/self-paced-courses/", "Day 3, 4, 8"),
        ("Generative AI Explained (NVIDIA DLI)", "https://www.nvidia.com/en-us/training/self-paced-courses/", "Day 3"),
        ("Augment Your LLM Using RAG (NVIDIA DLI)", "https://www.nvidia.com/en-us/training/self-paced-courses/", "Day 4, 8"),
        ("Cursor Get Started", "https://docs.cursor.com/get-started/introduction", "Day 4"),
        ("Cursor Agent mode", "https://docs.cursor.com/chat/agent", "Day 6+"),
        ("Git basics book", "https://git-scm.com/book/en/v2/Getting-Started-About-Version-Control", "Reference"),
        ("GitHub CLI", "https://cli.github.com/", "Appendix C if needed"),
    ]
    for label, url, when in links:
        add_link_line(doc, label, url, when)


def populate_appendix_c(doc: Document) -> None:
    add_appendix_opener(doc, "Appendix C", "Support and Troubleshooting", "Use when a step fails")
    add_display_heading(doc, "Push failed: Permission denied (403)", 2)
    add_numbered_step(doc, 1, "Install GitHub CLI", [])
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://cli.github.com/", "https://cli.github.com/")
    add_numbered_step(doc, 2, "Log in as your account", [], "gh auth login", "Browser login succeeds")
    add_numbered_step(doc, 3, "Clear saved password if needed", [], "cmdkey /delete:LegacyGeneric:target=git:https://github.com", "Credential deleted")
    add_numbered_step(doc, 4, "Retry push", [], "git push", "Success")
    add_display_heading(doc, "Git asks who you are", 2)
    add_body(doc, "Run git config name and email from Day 1.")
    add_display_heading(doc, "Push failed: remote hung up", 2)
    add_code_block(doc, "python scripts/download_oer.py")
    add_body(doc, "If still failing, log in progress/BLOCKERS.md and contact management.")


def populate_appendix_d(doc: Document) -> None:
    add_appendix_opener(doc, "Appendix D", "Onboarding Completion Criteria", "Review with management after Day 10")
    add_display_heading(doc, "Week 1 (Days 1 to 5)", 2)
    for item in [
        "Git workflow completed without assistance",
        "Session recordings uploaded each weekday",
        "UStud runs at http://localhost:5173",
        "CURRENT_STATE.md lists three product gaps",
        "At least one improvement pushed",
        "Session Start and Session End used in Cursor",
    ]:
        add_checklist_item(doc, item)
    add_display_heading(doc, "Week 2 (Days 6 to 10)", 2)
    for item in [
        "Three development tasks pushed to GitHub",
        "WEEK2_LOG.md completed for all five days",
        "DAILY_LOG.md includes two-week summary",
        "Offline tutor tested and documented",
        "Daily Session Start / Session End without prompts",
    ]:
        add_checklist_item(doc, item)


def populate_document(doc: Document) -> None:
    populate_intro(doc)
    populate_day_0(doc)
    populate_day_1(doc)
    populate_day_2(doc)
    populate_day_3(doc)
    populate_day_4(doc)
    populate_day_5(doc)
    populate_day_6(doc)
    populate_day_7(doc)
    populate_day_8(doc)
    populate_day_9(doc)
    populate_day_10(doc)
    populate_appendix_a(doc)
    populate_appendix_b(doc)
    populate_appendix_c(doc)
    populate_appendix_d(doc)

    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(24)
    fr = fp.add_run("AXIOM - Advanced - Intelligent - Practical - Grounded")
    _run_font(fr, FONT_MONO, 8, INK_MUTED)
