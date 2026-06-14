#!/usr/bin/env python3
"""Generate AXIOM-styled employee onboarding Word document (beginner-friendly)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "cursor" / "onboarding" / "AXIOM_EMPLOYEE_ONBOARDING.docx"

# AXIOM tokens
BLACK = RGBColor(0x08, 0x08, 0x08)
INK_PRIMARY = RGBColor(0xF0, 0xF0, 0xF0)
INK_SECONDARY = RGBColor(0xC4, 0xC4, 0xC4)
INK_MUTED = RGBColor(0x9A, 0x9A, 0x9A)
INK_FAINT = RGBColor(0x6E, 0x6E, 0x6E)
BODY_DARK = RGBColor(0x22, 0x22, 0x22)
BODY_MID = RGBColor(0x55, 0x55, 0x55)
COMMAND_LIVE = RGBColor(0x4A, 0x9E, 0xFF)
COMMAND_STABLE = RGBColor(0x3D, 0xD6, 0x8C)
COMMAND_WATCH = RGBColor(0xE8, 0xA8, 0x38)

FONT_SANS = "Calibri"
FONT_DISPLAY = "Arial"
FONT_MONO = "Consolas"


def _set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text: str, url: str, hex_color: str = "4A9EFF") -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), hex_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color_el)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _run_font(run, name: str, size_pt: float, color: RGBColor, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_spacer(doc: Document, pts: float = 12) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_section_divider(doc: Document, title: str, subtitle: str = "") -> None:
    """Full-width dark divider with optional page break for major sections."""
    add_page_break(doc)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "080808")
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:color"), "1A1A1A")
        borders.append(el)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title.upper())
    _run_font(r, FONT_MONO, 12, INK_PRIMARY, bold=True)
    if subtitle:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(12)
        r2 = p2.add_run(subtitle)
        _run_font(r2, FONT_SANS, 10.5, INK_SECONDARY)
    add_spacer(doc, 18)


def add_section_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    _run_font(run, FONT_MONO, 9, INK_MUTED)
    run.font.all_caps = True


def add_part_heading(
    doc: Document,
    part_num: int,
    title: str,
    subtitle: str = "",
    *,
    page_break: bool = True,
) -> None:
    if page_break:
        add_page_break(doc)
    add_spacer(doc, 8)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"PART {part_num}")
    _run_font(r, FONT_MONO, 10, COMMAND_LIVE, bold=True)
    add_spacer(doc, 4)
    hp = doc.add_paragraph()
    hp.paragraph_format.space_after = Pt(6)
    hr = hp.add_run(title)
    _run_font(hr, FONT_DISPLAY, 20, BLACK, bold=True)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(14)
        sr = sp.add_run(subtitle)
        _run_font(sr, FONT_SANS, 11, BODY_MID)


def add_display_heading(doc: Document, text: str, level: int = 2) -> None:
    sizes = {2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    _run_font(run, FONT_DISPLAY, sizes.get(level, 12), BLACK, bold=True)


def add_body(doc: Document, text: str, space_after: float = 10) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    _run_font(run, FONT_SANS, 11, BODY_DARK)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    _run_font(run, FONT_SANS, 11, BODY_DARK)


def add_numbered_step(
    doc: Document,
    number: int,
    title: str,
    instructions: list[str],
    command: str | None = None,
    success: str | None = None,
) -> None:
    """One clear step: title, plain instructions, optional single command, optional success check."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"Step {number}. {title}")
    _run_font(r, FONT_SANS, 11.5, BLACK, bold=True)

    for line in instructions:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(4)
        bp.paragraph_format.left_indent = Inches(0.25)
        bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        bp.paragraph_format.line_spacing = 1.25
        br = bp.add_run(line)
        _run_font(br, FONT_SANS, 11, BODY_DARK)

    if command:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(6)
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.left_indent = Inches(0.25)
        cr = cp.add_run("Copy this line, paste into PowerShell, press Enter:")
        _run_font(cr, FONT_SANS, 10, INK_MUTED, italic=True)
        add_code_block(doc, command, indent=True)

    if success:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        sp.paragraph_format.space_after = Pt(10)
        sp.paragraph_format.left_indent = Inches(0.25)
        sr1 = sp.add_run("You should see: ")
        _run_font(sr1, FONT_SANS, 10.5, COMMAND_STABLE, bold=True)
        sr2 = sp.add_run(success)
        _run_font(sr2, FONT_SANS, 10.5, BODY_DARK)


def add_callout(doc: Document, label: str, text: str, kind: str = "tip") -> None:
    colors = {"tip": "E8F4FF", "watch": "FFF8E8", "stable": "E8F8EF"}
    fill = colors.get(kind, colors["tip"])
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "CCCCCC")
        borders.append(el)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run(f"{label.upper()}  ")
    color = COMMAND_WATCH if kind == "watch" else COMMAND_LIVE if kind == "tip" else COMMAND_STABLE
    _run_font(r1, FONT_MONO, 8.5, color, bold=True)
    r2 = p.add_run(text)
    _run_font(r2, FONT_SANS, 10.5, BODY_DARK)
    add_spacer(doc, 8)


def add_code_block(doc: Document, code: str, indent: bool = False) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, "0A0A0A")
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "1A1A1A")
        borders.append(el)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code.strip())
    _run_font(run, FONT_MONO, 10, INK_PRIMARY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    if indent:
        spacer.paragraph_format.left_indent = Inches(0.25)


def add_link_line(doc: Document, label: str, url: str, when: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(f"• {label}")
    _run_font(run, FONT_SANS, 11, BODY_DARK, bold=True)
    if when:
        run2 = p.add_run(f"  ({when})  ")
        _run_font(run2, FONT_SANS, 10, INK_MUTED)
    add_hyperlink(p, url, url)


def add_checklist_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(f"[ ]  {text}")
    _run_font(r, FONT_SANS, 11, BODY_DARK)


def add_toc_item(doc: Document, part: str, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(f"{part}  ")
    _run_font(r1, FONT_MONO, 10, COMMAND_LIVE, bold=True)
    r2 = p.add_run(title)
    _run_font(r2, FONT_SANS, 11, BODY_DARK)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Cover
    header_table = doc.add_table(rows=1, cols=1)
    header_cell = header_table.rows[0].cells[0]
    _set_cell_shading(header_cell, "080808")
    hp = header_cell.paragraphs[0]
    hr1 = hp.add_run("AXIOM  |  EMPLOYEE ONBOARDING\n")
    _run_font(hr1, FONT_MONO, 9, INK_MUTED)
    hr2 = hp.add_run("Two-Week Onboarding Program\n")
    _run_font(hr2, FONT_DISPLAY, 22, INK_PRIMARY, bold=True)
    hr3 = hp.add_run("Week 1: Foundations  |  Week 2: Applied Development")
    _run_font(hr3, FONT_SANS, 11, INK_SECONDARY)

    add_spacer(doc, 16)
    add_body(
        doc,
        "This guide assumes no prior experience with GitHub, the terminal, or development tools. "
        "Follow each part in order. Each step states what to click, what to copy, and what you should see on screen.",
        space_after=14,
    )

    add_callout(
        doc,
        "Schedule",
        "Day 0 (laptop setup): 2 to 3 hours. Weeks 1 and 2: about 1 hour per day, Monday through Friday. "
        "Session recordings are required during Week 1 only unless management requests otherwise.",
        "tip",
    )

    add_page_break(doc)

    # --- PART 0: READ FIRST ---
    add_part_heading(
        doc,
        0,
        "Read This Before Anything Else",
        "Required reading before Day 0.",
        page_break=False,
    )

    add_display_heading(doc, "What is PowerShell?", 3)
    add_body(
        doc,
        "PowerShell is a black or blue window where you type commands instead of clicking. "
        "We use it to talk to GitHub and run the project. You will copy commands from this document and paste them there.",
    )
    add_numbered_step(
        doc,
        1,
        "Open PowerShell",
        [
            "Click the Windows Start button (bottom left).",
            "Type: PowerShell",
            "Click Windows PowerShell (not PowerShell ISE).",
        ],
        success="A window opens with text like PS C:\\Users\\YourName>",
    )

    add_display_heading(doc, "How to copy and paste a command", 3)
    add_numbered_step(
        doc,
        1,
        "Copy from this document",
        ["Select the text inside the dark box.", "Press Ctrl+C on your keyboard."],
    )
    add_numbered_step(
        doc,
        2,
        "Paste into PowerShell",
        [
            "Click inside the PowerShell window.",
            "Right-click to paste (or press Ctrl+V).",
            "Press Enter on your keyboard.",
        ],
        success="The command runs. New text appears. If you see red error text, read the step again or see Part 10 (Support).",
    )

    add_callout(
        doc,
        "Important",
        "Run one command at a time unless this guide shows several lines in one box labeled run all. "
        "Wait for each command to finish before running the next.",
        "watch",
    )

    add_display_heading(doc, "What you need from management", 3)
    add_checklist_item(doc, "An email invite to join the GitHub repo axiom-forhumanity/ustud (accept before Monday)")
    add_checklist_item(doc, "A link to upload your daily screen recordings (Google Drive or similar)")
    add_checklist_item(doc, "A contact email if you get stuck")

    add_display_heading(doc, "Table of contents", 3)
    add_toc_item(doc, "Part 1", "Day 0 — Install software")
    add_toc_item(doc, "Part 2", "Day 0 — Create accounts")
    add_toc_item(doc, "Part 3", "GitHub — concepts and setup")
    add_toc_item(doc, "Part 4", "Week 1, Monday — Download the project")
    add_toc_item(doc, "Part 5", "Week 1, Tuesday–Friday — Daily plan")
    add_toc_item(doc, "Part 6", "Week 1 — Session recordings")
    add_toc_item(doc, "Part 7", "Cursor — session commands")
    add_toc_item(doc, "Part 8", "Week 2, Monday–Friday — Development plan")
    add_toc_item(doc, "Part 9", "Learning resources")
    add_toc_item(doc, "Part 10", "Support and troubleshooting")
    add_toc_item(doc, "Part 11", "Onboarding completion criteria")

    add_section_divider(doc, "Week 1 — Foundations", "Setup, GitHub, automation, AI concepts, first run of UStud")

    # --- PART 1: DAY 0 INSTALLS ---
    add_part_heading(
        doc,
        1,
        "Day 0 — Install Software",
        "Complete when your laptop arrives, before Week 1 Monday. Windows 10 or 11 required.",
        page_break=False,
    )

    add_callout(doc, "Order matters", "Install programs in the order below. After each install, run the check command if one is shown.", "tip")

    installs = [
        (
            "Install Git",
            "Git is the tool that saves and sends your work to GitHub.",
            "https://git-scm.com/download/win",
            [
                "Open the download link. Download runs automatically.",
                "Run the installer. Click Next on every screen (defaults are fine).",
                "When finished, open PowerShell.",
            ],
            "git --version",
            "git version 2.x.x (any number is fine)",
        ),
        (
            "Install Python",
            "Python runs the backend of the UStud app.",
            "https://www.python.org/downloads/",
            [
                "Open the download link. Click the yellow Download button.",
                "Run the installer.",
                "On the FIRST screen, check the box: Add python.exe to PATH. This is important.",
                "Click Install Now. Wait until it finishes.",
            ],
            "python --version",
            "Python 3.10 or higher",
        ),
        (
            "Install Node.js",
            "Node.js runs the user interface of the app.",
            "https://nodejs.org/",
            [
                "Open the download link.",
                "Download the LTS version (recommended).",
                "Run the installer. Click Next through all screens.",
            ],
            "node --version",
            "A version number like v20.x.x",
        ),
        (
            "Install Cursor",
            "Cursor is where you write code and talk to the AI assistant.",
            "https://cursor.com/download",
            [
                "Open the download link. Download and install.",
                "Open Cursor from the Start menu.",
                "Create a free account or sign in when asked.",
            ],
            None,
            "Cursor opens and you are signed in",
        ),
        (
            "Install Ollama",
            "Ollama runs the offline AI tutor on your laptop.",
            "https://ollama.com/download",
            [
                "Open the download link. Download for Windows.",
                "Run the installer (admin permission may be required).",
            ],
            "ollama --version",
            "A version number",
        ),
        (
            "Install OBS Studio",
            "OBS records your screen and webcam during Week 1 sessions.",
            "https://obsproject.com/",
            [
                "Open the download link. Download and install.",
                "Open OBS once to confirm it launches.",
                "Full OBS setup is in Part 6.",
            ],
            None,
            "OBS opens without crashing",
        ),
    ]

    step_num = 1
    for title, why, url, click_steps, verify_cmd, verify_ok in installs:
        add_display_heading(doc, title, 2)
        add_body(doc, why)
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(6)
        lr = lp.add_run("Download link: ")
        _run_font(lr, FONT_SANS, 11, BODY_DARK, bold=True)
        add_hyperlink(lp, url, url)
        for click in click_steps:
            add_numbered_step(doc, step_num, click, [])
            step_num += 1
        if verify_cmd:
            add_numbered_step(
                doc,
                step_num,
                "Check that it worked",
                ["Open PowerShell. Paste and run:"],
                command=verify_cmd,
                success=verify_ok,
            )
            step_num += 1
        add_spacer(doc, 10)

    # --- PART 2: ACCOUNTS ---
    add_part_heading(doc, 2, "Day 0 — Create Online Accounts", "Two free accounts. Do this after installing software.")

    add_display_heading(doc, "Create a GitHub account", 2)
    add_body(doc, "GitHub is the website where the team project lives. Management needs your username to give you access.")
    add_numbered_step(doc, 1, "Go to the signup page", [], command=None)
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://github.com/signup", "https://github.com/signup")
    add_numbered_step(
        doc,
        2,
        "Create your account",
        [
            "Pick a username (write it down).",
            "Use your real email. Verify your email when GitHub sends a message.",
            "Email your username to management so they can add you to the project.",
        ],
    )
    add_numbered_step(
        doc,
        3,
        "Accept the team invite",
        [
            "Check your email for an invitation to axiom-forhumanity/ustud.",
            "Click Accept invitation.",
            "Open this link in your browser to confirm you see the project:",
        ],
    )
    lp2 = doc.add_paragraph()
    add_hyperlink(lp2, "https://github.com/axiom-forhumanity/ustud", "https://github.com/axiom-forhumanity/ustud")
    add_body(doc, "You should see files and folders, not a 404 error page.", space_after=14)

    add_display_heading(doc, "Create an NVIDIA account (for free AI courses)", 2)
    add_body(doc, "You will use this on Wednesday and Thursday for free training videos.")
    lp3 = doc.add_paragraph()
    add_hyperlink(lp3, "https://developer.nvidia.com/login", "https://developer.nvidia.com/login")
    add_numbered_step(doc, 1, "Sign up for a free account", ["Follow the website steps. Verify your email if asked."])

    # --- PART 3: GITHUB EXPLAINED ---
    add_part_heading(
        doc,
        3,
        "GitHub — Concepts and Setup",
        "Read on Week 1 Monday before you download the project.",
    )

    add_body(
        doc,
        "Think of the project like a shared folder in the cloud. Your laptop has a copy. "
        "When you save work, you send it back to the cloud so management can see it.",
        space_after=12,
    )

    terms = [
        ("Repository (repo)", "The whole project: all files and all history of changes."),
        ("Clone", "Download the project from GitHub to your laptop (you do this once)."),
        ("Pull", "Get the newest changes from GitHub before you start working each day."),
        ("Commit", "Save a snapshot of your changes on your laptop with a short note."),
        ("Push", "Send your saved snapshots up to GitHub so the team can see them."),
    ]
    for word, meaning in terms:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r1 = p.add_run(f"{word}: ")
        _run_font(r1, FONT_SANS, 11, BLACK, bold=True)
        r2 = p.add_run(meaning)
        _run_font(r2, FONT_SANS, 11, BODY_DARK)

    add_callout(
        doc,
        "Daily habit",
        "Every work day: pull first (get updates), do your work, then commit and push (send your work up).",
        "stable",
    )

    add_display_heading(doc, "Tell Git your name (one time only)", 2)
    add_body(doc, "Replace the name and email with yours. Run these two lines in PowerShell:")
    add_code_block(doc, 'git config --global user.name "Your Full Name"')
    add_code_block(doc, 'git config --global user.email "your.email@example.com"')

    # --- PART 4: WEEK 1 MONDAY ---
    add_part_heading(
        doc,
        4,
        "Week 1, Monday — Download the Project",
        "About 1 hour. Start your session recording first (see Part 6).",
    )

    add_callout(doc, "Before you start", "Start OBS recording (screen + webcam). Pause OBS if you type a password.", "watch")

    monday_steps = [
        (
            "Open PowerShell",
            ["Use the Start menu. Search PowerShell. Open it."],
            None,
            "Window is open",
        ),
        (
            "Go to your Documents folder",
            ["This is where the project will live."],
            "cd $HOME\\Documents",
            "The prompt may change to show Documents",
        ),
        (
            "Download the project from GitHub (clone)",
            ["This may take 1 to 2 minutes. Wait until it finishes."],
            "git clone https://github.com/axiom-forhumanity/ustud.git",
            "Text like Cloning into ustud... then it stops with no red errors",
        ),
        (
            "Enter the project folder",
            ["All future commands assume you are inside this folder."],
            "cd ustud",
            "Prompt shows ...\\Documents\\ustud",
        ),
        (
            "Open the project in Cursor",
            [
                "Open Cursor.",
                "Click File, then Open Folder.",
                "Go to Documents, then ustud, then click Select Folder.",
            ],
            None,
            "You see files on the left side: README.md, cursor, docs, etc.",
        ),
        (
            "Open your Week 1 log file",
            [
                "In Cursor's left file list, open: progress, then WEEK1_LOG.md",
                "Fill in the Monday section: what you learned, link to your recording (after upload), one sentence takeaway.",
                "Save the file (Ctrl+S).",
            ],
            None,
            "Monday section has your notes",
        ),
        (
            "Send your changes to GitHub — step A: get latest",
            ["Always pull before you push."],
            "git pull",
            "Already up to date or a success message",
        ),
        (
            "Send your changes — step B: stage files",
            ["This tells Git which files you changed."],
            "git add .",
            "No output is normal",
        ),
        (
            "Send your changes — step C: save with a message",
            ["The message describes what you did today."],
            'git commit -m "Week1 Day1: completed GitHub setup and first push"',
            "A line like 1 file changed",
        ),
        (
            "Send your changes — step D: upload to GitHub",
            ["This sends your work to the team repo."],
            "git push",
            "Progress lines ending without red errors",
        ),
    ]

    for i, (title, instr, cmd, ok) in enumerate(monday_steps, 1):
        add_numbered_step(doc, i, title, instr, command=cmd, success=ok)

    add_body(doc, "Optional reading (30 min): GitHub Hello World tutorial — good for beginners.")
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://docs.github.com/en/get-started/start-your-journey/hello-world", "https://docs.github.com/en/get-started/start-your-journey/hello-world")

    # --- PART 5: WEEK 1 TUE-FRI ---
    add_part_heading(
        doc,
        5,
        "Week 1, Tuesday–Friday — Daily Plan",
        "Same order each day: start recording, complete the lesson, upload video, update log, push to GitHub.",
    )

    add_callout(
        doc,
        "Every day starts the same",
        "1. Start OBS.  2. Open PowerShell.  3. cd $HOME\\Documents\\ustud  4. git pull",
        "stable",
    )

    # Tuesday
    add_display_heading(doc, "Week 1, Tuesday — Automation", 2)
    add_body(doc, "Automation means the computer does steps for you. Example: setup.ps1 installs packages automatically instead of you clicking many times.")
    add_numbered_step(doc, 1, "Pull latest code", ["In PowerShell, from the ustud folder:"], "git pull", "Success message")
    add_numbered_step(
        doc,
        2,
        "Run the setup script",
        ["Still in the ustud folder. This installs Python packages and downloads the AI model. May take several minutes."],
        ".\\setup.ps1",
        "Green success messages. If Ollama is missing, note it in progress/BLOCKERS.md",
    )
    add_numbered_step(
        doc,
        3,
        "Update and push your log",
        ["Edit progress/WEEK1_LOG.md (Tuesday section). Then run git add ., commit, push (same as Monday steps 7–10)."],
    )
    add_body(doc, "Optional reading: GitHub Actions quickstart (20 min).")
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://docs.github.com/en/actions/quickstart", "https://docs.github.com/en/actions/quickstart")

    # Wednesday
    add_page_break(doc)
    add_display_heading(doc, "Week 1, Wednesday — AI and LLMs", 2)
    add_body(doc, "An LLM (Large Language Model) is AI that reads and writes text. UStud uses a local LLM so students can learn without internet.")
    add_numbered_step(doc, 1, "Open NVIDIA free courses", ["Sign in with your NVIDIA account from Part 2."])
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://www.nvidia.com/en-us/training/self-paced-courses/", "https://www.nvidia.com/en-us/training/self-paced-courses/")
    add_numbered_step(
        doc,
        2,
        "Enroll in Generative AI Explained",
        ["Use the Free Courses filter.", "Do the first hour today. Finish on Thursday if needed."],
    )
    add_numbered_step(
        doc,
        3,
        "Write three short answers in WEEK1_LOG.md (Wednesday section)",
        [
            "What is an LLM? (your own words)",
            "What is offline AI?",
            "Why does UStud use local models?",
            "Then git add ., commit, push.",
        ],
    )

    # Thursday
    add_page_break(doc)
    add_display_heading(doc, "Week 1, Thursday — Agents and Cursor", 2)
    add_body(doc, "A chatbot only talks. An agent can also read files, edit code, and run commands. Cursor in Agent mode is an agent.")
    add_numbered_step(doc, 1, "Continue NVIDIA course or start Augment Your LLM Using RAG", ["Same NVIDIA link as Wednesday."])
    add_numbered_step(
        doc,
        2,
        "Start a Cursor session",
        [
            "Open Cursor with the ustud folder.",
            "Open the Chat panel.",
            "Copy the entire Session Start block from Part 7 below and paste into chat. Press Enter.",
            "Wait for Cursor to respond.",
        ],
    )
    add_numbered_step(
        doc,
        3,
        "Ask Cursor a question",
        ["Type or paste: Explain the folder structure of this repo in simple terms."],
    )
    add_numbered_step(
        doc,
        4,
        "End the session",
        ["Copy the Session End block from Part 7. Paste into chat.", "Review what Cursor wrote in the progress files.", "git add ., commit, push."],
    )
    add_body(doc, "Optional: Cursor Get Started docs.")
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://docs.cursor.com/get-started/introduction", "https://docs.cursor.com/get-started/introduction")

    # Friday
    add_page_break(doc)
    add_display_heading(doc, "Week 1, Friday — Run the Application", 2)
    add_body(doc, "Run UStud on your laptop and deliver your first small improvement.")
    add_numbered_step(doc, 1, "Pull latest", [], "git pull", "Success")
    add_numbered_step(doc, 2, "Install dependencies (first time only)", ["Run each line, wait for each to finish:"], "npm install", "Completes without errors")
    add_numbered_step(doc, 3, "Install UI dependencies", [], "cd boiler", "Folder changes")
    add_numbered_step(doc, 4, "Continue UI install", [], "npm install", "Completes")
    add_numbered_step(doc, 5, "Go back to project root", [], "cd ..", "Back in ustud folder")
    add_numbered_step(
        doc,
        6,
        "Start the app",
        ["Leave this window open while the app runs."],
        "npm run dev:all",
        "Messages about API and UI starting",
    )
    add_numbered_step(
        doc,
        7,
        "Open the app in your browser",
        ["Open Chrome or Edge.", "Go to: http://localhost:5173", "You should see the UStud interface."],
        None,
        "App loads in the browser",
    )
    add_numbered_step(
        doc,
        8,
        "Document and improve",
        [
            "With Cursor, update progress/CURRENT_STATE.md: top 3 gaps, what works, what does not.",
            "Make one small fix (typo, label, tiny bug).",
            "Session End from Part 7, then git add ., commit, push.",
        ],
    )

    add_callout(
        doc,
        "End of Week 1",
        "Confirm: GitHub push works, session recordings uploaded Mon–Fri, UStud runs at http://localhost:5173, "
        "and progress/CURRENT_STATE.md lists three product gaps.",
        "stable",
    )

    # --- PART 6: OBS ---
    add_part_heading(
        doc,
        6,
        "Week 1 — Session Recordings",
        "Required Monday through Friday during Week 1, about 1 hour per session.",
    )

    add_body(
        doc,
        "Management will send you a folder link for uploads. Record your screen and your face. "
        "Upload the same day. Pause OBS when typing passwords.",
        space_after=12,
    )

    obs_steps = [
        ("Open OBS Studio", ["From the Start menu."], None),
        ("Create a scene", ["Bottom left: click + under Scenes.", "Name it Week1.", "Press Enter."], None),
        ("Add your screen", ["Under Sources, click +.", "Choose Display Capture, click OK.", "Select your monitor, click OK."], None),
        ("Add your webcam", ["Under Sources, click + again.", "Choose Video Capture Device, click OK.", "Select your camera.", "Drag the webcam box to a corner. Make it small."], None),
        ("Set video quality (optional)", ["File menu, Settings, Video.", "If uploads are slow, set Base and Output to 1280x720."], None),
        ("Test recording", ["Click Start Recording.", "Wait 30 seconds.", "Click Stop Recording.", "File, Show Recordings, play the file."], "You see your screen and your face"),
        ("Each work day", ["Click Start Recording before you begin.", "Work for about 1 hour.", "Click Stop Recording.", "Name file: Week1_Day1_YYYY-MM-DD_YourName.mp4 (Day1=Mon, Day5=Fri)."], None),
        ("Upload", ["Open the folder link from management.", "Upload the MP4.", "Copy the file link into progress/WEEK1_LOG.md under Trust recording URL.", "git add ., commit, push."], None),
    ]
    for i, (title, instr, ok) in enumerate(obs_steps, 1):
        add_numbered_step(doc, i, title, instr, success=ok)

    # --- PART 7: CURSOR COMMANDS ---
    add_part_heading(
        doc,
        7,
        "Cursor — Session Commands",
        "Use at the start and end of every work session during Week 2 and ongoing development.",
    )

    add_display_heading(doc, "Session Start — paste at the beginning", 2)
    add_body(doc, "Copy everything in the box below into Cursor chat:")
    add_code_block(
        doc,
        """Read and follow cursor/AXIOM_OPERATING_INSTRUCTIONS.md completely.

Also read:
- docs/PROJECT_BRIEF.md
- progress/CURRENT_STATE.md
- progress/BLOCKERS.md
- the latest entry in progress/DAILY_LOG.md

Confirm you understand the progress logging requirements.

Then tell me:
1. Current project state
2. Open blockers
3. What I should focus on today

Do not start coding until you have read these files.""",
    )

    add_display_heading(doc, "Session End — paste before git push", 2)
    add_body(doc, "Copy everything in the box below into Cursor chat:")
    add_code_block(
        doc,
        """Session ending. Follow cursor/AXIOM_OPERATING_INSTRUCTIONS.md session END workflow.

Update:
- progress/DAILY_LOG.md (append today's entry)
- progress/CURRENT_STATE.md
- progress/BLOCKERS.md (if needed)
- docs/DECISION_LOG.md (if any decisions were made)

Then show me a summary of exactly what you logged so I can review before I commit and push to GitHub.""",
    )

    add_section_divider(doc, "Week 2 — Applied Development", "Build on UStud, deliver improvements, establish daily workflow")

    # --- PART 8: WEEK 2 ---
    add_part_heading(
        doc,
        8,
        "Week 2, Monday–Friday — Development Plan",
        "About 1 hour per day. Session recordings are not required unless management asks. Daily git push is required.",
        page_break=False,
    )

    add_callout(
        doc,
        "Daily workflow",
        "1. Open PowerShell.  2. cd $HOME\\Documents\\ustud  3. git pull  "
        "4. Session Start in Cursor  5. Complete that day's task  6. Session End  7. git add ., commit, push",
        "stable",
    )

    add_body(doc, "Update progress/WEEK2_LOG.md each day (same format as WEEK1_LOG.md). Use progress/DAILY_LOG.md for detailed work entries.")

    # Week 2 Monday
    add_display_heading(doc, "Week 2, Monday — Codebase Orientation", 2)
    add_body(doc, "Objective: Understand where code, UI, and course content live before making changes.")
    add_numbered_step(doc, 1, "Pull latest code", ["Open PowerShell in the ustud folder."], "git pull", "Success message")
    add_numbered_step(
        doc,
        2,
        "Start Cursor session",
        ["Paste Session Start from Part 7.", "Ask Cursor to walk through these folders one at a time:"],
    )
    add_bullet(doc, "src/ustud/ - Python backend (API, AI tutor logic)")
    add_bullet(doc, "boiler/ - React user interface")
    add_bullet(doc, "modules/ - course content and curriculum")
    add_bullet(doc, "progress/ - your logs and project status files")
    add_numbered_step(
        doc,
        3,
        "Update progress/CURRENT_STATE.md",
        [
            "Expand your Week 1 gap list into full sentences.",
            "Add a section: Recommended next task (pick the smallest gap).",
            "Session End, then git add ., commit, push.",
        ],
    )
    add_numbered_step(
        doc,
        4,
        "Log Week 2 progress",
        ["Fill Monday section in progress/WEEK2_LOG.md.", "Commit and push."],
    )

    add_page_break(doc)
    add_display_heading(doc, "Week 2, Tuesday — First Development Task", 2)
    add_body(doc, "Objective: Ship one improvement tied to gap #1 in CURRENT_STATE.md.")
    add_numbered_step(doc, 1, "Pull and start session", [], "git pull", "Success")
    add_numbered_step(
        doc,
        2,
        "Implement gap #1",
        [
            "Session Start in Cursor.",
            "Tell Cursor: Help me implement the first gap listed in progress/CURRENT_STATE.md. Keep the change small.",
            "Test locally if the app is running (npm run dev:all).",
        ],
    )
    add_numbered_step(
        doc,
        3,
        "Document and push",
        [
            "Session End — confirm DAILY_LOG.md and CURRENT_STATE.md are updated.",
            "Fill Tuesday section in WEEK2_LOG.md.",
        ],
    )
    add_code_block(doc, 'git add .\ngit commit -m "Week2 Day2: first development task - [brief description]"\ngit push')

    add_page_break(doc)
    add_display_heading(doc, "Week 2, Wednesday — AI Integration", 2)
    add_body(doc, "Objective: Connect training to the product. Finish NVIDIA RAG content and test the offline tutor.")
    add_numbered_step(doc, 1, "Complete NVIDIA course", ["Finish Generative AI Explained or complete Augment Your LLM Using RAG."])
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://www.nvidia.com/en-us/training/self-paced-courses/", "https://www.nvidia.com/en-us/training/self-paced-courses/")
    add_numbered_step(
        doc,
        2,
        "Test UStud tutor",
        [
            "Ensure the app is running (npm run dev:all).",
            "Open http://localhost:5173",
            "Ask the offline tutor one question about a lesson topic.",
            "Note in WEEK2_LOG.md: Did it work? What would improve the response?",
        ],
    )
    add_numbered_step(doc, 3, "Push your notes", [], "git pull", "Success")
    add_numbered_step(
        doc,
        4,
        "Commit log updates",
        ["After editing WEEK2_LOG.md and DAILY_LOG.md:"],
        'git add .\ngit commit -m "Week2 Day3: AI integration notes and tutor test"\ngit push',
    )

    add_page_break(doc)
    add_display_heading(doc, "Week 2, Thursday — Second Development Task", 2)
    add_body(doc, "Objective: Deliver one UI or module improvement (gap #2 or related polish).")
    add_numbered_step(doc, 1, "Pull and Session Start", [], "git pull", "Success")
    add_numbered_step(
        doc,
        2,
        "Build and test",
        [
            "Work in boiler/ (UI) or modules/ (content) with Cursor.",
            "Verify the change in the browser or app.",
            "Session End before push.",
        ],
    )
    add_code_block(doc, 'git add .\ngit commit -m "Week2 Day4: second development task - [brief description]"\ngit push')
    add_numbered_step(
        doc,
        3,
        "Update WEEK2_LOG.md",
        ["Record what you changed and whether it is working."],
    )

    add_page_break(doc)
    add_display_heading(doc, "Week 2, Friday — Third Deliverable and Summary", 2)
    add_body(doc, "Objective: Complete a third small improvement and submit your two-week onboarding summary.")
    add_numbered_step(doc, 1, "Pull latest", [], "git pull", "Success")
    add_numbered_step(
        doc,
        2,
        "Third improvement",
        [
            "Address gap #3, or refine one of this week's changes.",
            "Confirm UStud still runs locally.",
        ],
    )
    add_numbered_step(
        doc,
        3,
        "Write two-week summary",
        [
            "In progress/DAILY_LOG.md, add a summary entry covering:",
            "What you built in Week 2",
            "What remains blocked or unclear",
            "Your recommended focus for Week 3",
            "Complete Friday section in WEEK2_LOG.md.",
            "Session End, then push.",
        ],
    )
    add_code_block(doc, 'git add .\ngit commit -m "Week2 Day5: third deliverable and onboarding summary"\ngit push')

    add_callout(
        doc,
        "End of Week 2",
        "Review Part 11 completion criteria with management. Ongoing daily work follows cursor/EMPLOYEE_ONBOARDING.md in the repo.",
        "stable",
    )

    # --- PART 9: LEARNING LINKS ---
    add_part_heading(doc, 9, "Learning Resources", "Hyperlinks for courses and reference material.")

    links = [
        ("Team project on GitHub", "https://github.com/axiom-forhumanity/ustud", "Always"),
        ("Project brief", "https://github.com/axiom-forhumanity/ustud/blob/main/docs/PROJECT_BRIEF.md", "Week 2"),
        ("AXIOM vision", "https://github.com/axiom-forhumanity/ustud/blob/main/docs/AXIOM_VISION.md", "Week 2"),
        ("GitHub Hello World", "https://docs.github.com/en/get-started/start-your-journey/hello-world", "Week 1 Mon"),
        ("GitHub Actions intro", "https://docs.github.com/en/actions/quickstart", "Week 1 Tue"),
        ("NVIDIA free courses", "https://www.nvidia.com/en-us/training/self-paced-courses/", "Week 1 Wed–Thu, Week 2 Wed"),
        ("Cursor introduction", "https://docs.cursor.com/get-started/introduction", "Week 1 Thu"),
        ("Cursor Agent mode", "https://docs.cursor.com/chat/agent", "Week 2"),
        ("Git basics book", "https://git-scm.com/book/en/v2/Getting-Started-About-Version-Control", "Reference"),
    ]
    for label, url, when in links:
        add_link_line(doc, label, url, when)

    add_body(doc, "Save NVIDIA course certificates as screenshots in your Week 1 upload folder if provided.")

    # --- PART 10: TROUBLESHOOTING ---
    add_part_heading(doc, 10, "Support and Troubleshooting", "Common issues and resolution steps.")

    add_display_heading(doc, "Push failed: Permission denied (403)", 2)
    add_body(doc, "You are logged into the wrong GitHub account.")
    add_numbered_step(doc, 1, "Install GitHub CLI", ["Download from:"])
    lp = doc.add_paragraph()
    add_hyperlink(lp, "https://cli.github.com/", "https://cli.github.com/")
    add_numbered_step(doc, 2, "Log in with YOUR account", [], "gh auth login", "Follow browser steps. Log in as the account management invited.")
    add_numbered_step(
        doc,
        3,
        "Clear old password (if still failing)",
        [],
        "cmdkey /delete:LegacyGeneric:target=git:https://github.com",
        "Credential deleted successfully",
    )
    add_numbered_step(doc, 4, "Try git push again", [], "git push", "Upload succeeds")

    add_display_heading(doc, "Git asks who you are", 2)
    add_body(doc, "Run the name and email commands from Part 3 again.")

    add_display_heading(doc, "Push failed: remote hung up", 2)
    add_body(doc, "Large files may cause this. Textbooks are not in Git. Download them separately:")
    add_code_block(doc, "python scripts/download_oer.py")
    add_body(doc, "If it still fails, write in progress/BLOCKERS.md and email management.")

    add_display_heading(doc, "Other issues", 2)
    add_body(doc, "1. Write what happened in progress/BLOCKERS.md.  2. git push.  3. Email management.")

    # --- PART 11: COMPLETION ---
    add_part_heading(doc, 11, "Onboarding Completion Criteria", "Confirm each item with management at the end of Week 2.")

    add_display_heading(doc, "Week 1 — Foundation", 2)
    for item in [
        "Git pull, git add, git commit, and git push completed without assistance",
        "Session recordings uploaded each weekday during Week 1",
        "UStud runs locally at http://localhost:5173",
        "progress/CURRENT_STATE.md lists three product gaps",
        "At least one small improvement pushed to GitHub",
        "Session Start and Session End used in Cursor",
    ]:
        add_checklist_item(doc, item)

    add_display_heading(doc, "Week 2 — Development", 2)
    for item in [
        "Three development tasks pushed to GitHub (Tue, Thu, Fri minimum)",
        "progress/WEEK2_LOG.md completed for all five weekdays",
        "progress/DAILY_LOG.md includes a two-week summary entry",
        "Offline tutor tested and results documented",
        "Daily Session Start / Session End workflow followed without prompts",
    ]:
        add_checklist_item(doc, item)

    add_spacer(doc, 12)
    add_body(doc, "After onboarding, your daily operating guide is cursor/EMPLOYEE_ONBOARDING.md in the repo.")
    add_body(doc, "Questions and blockers: progress/BLOCKERS.md or contact management directly.")

    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(24)
    fr = fp.add_run("AXIOM · Advanced · Intelligent · Practical · Grounded")
    _run_font(fr, FONT_MONO, 8, INK_FAINT)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
