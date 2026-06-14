#!/usr/bin/env python3
"""Generate AXIOM-styled USTUD project brief one-pager (Word)."""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

from onboarding_doc_common import (
    BODY_DARK,
    BODY_MID,
    BLACK,
    COMMAND_LIVE,
    COMMAND_STABLE,
    COMMAND_WATCH,
    FONT_DISPLAY,
    FONT_MONO,
    FONT_SANS,
    INK_MUTED,
    INK_PRIMARY,
    INK_SECONDARY,
    _run_font,
    _set_cell_shading,
)

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "USTUD_PROJECT_BRIEF_ONE_PAGER.docx"


def _label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    _run_font(r, FONT_MONO, 8, INK_MUTED)
    r.font.all_caps = True


def _body(doc: Document, text: str, size: float = 9.5, after: float = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    _run_font(r, FONT_SANS, size, BODY_DARK)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    _run_font(r, FONT_SANS, 9, BODY_DARK)


def _callout(doc: Document, tag: str, text: str, kind: str = "stable") -> None:
    fills = {"stable": "E8F8EF", "watch": "FFF8E8", "live": "E8F4FF"}
    colors = {"stable": COMMAND_STABLE, "watch": COMMAND_WATCH, "live": COMMAND_LIVE}
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, fills.get(kind, fills["stable"]))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(f"{tag.upper()}  ")
    _run_font(r1, FONT_MONO, 7.5, colors.get(kind, COMMAND_STABLE), bold=True)
    r2 = p.add_run(text)
    _run_font(r2, FONT_SANS, 9, BODY_DARK)


def build() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.5)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)

    # Header
    ht = doc.add_table(rows=1, cols=1)
    hc = ht.rows[0].cells[0]
    _set_cell_shading(hc, "080808")
    hp = hc.paragraphs[0]
    hp.paragraph_format.space_before = Pt(10)
    hp.paragraph_format.space_after = Pt(4)
    r0 = hp.add_run("AXIOM  |  PROJECT BRIEF\n")
    _run_font(r0, FONT_MONO, 8, INK_MUTED)
    r1 = hp.add_run("USTUD\n")
    _run_font(r1, FONT_DISPLAY, 20, INK_PRIMARY, bold=True)
    r2 = hp.add_run("Offline AI Education System")
    _run_font(r2, FONT_SANS, 10, INK_SECONDARY)

    _body(
        doc,
        "USTUD is an education technology project with a simple but powerful purpose: to make high-quality "
        "knowledge and tutoring available even when the internet is unavailable, unreliable, expensive, or limited.",
        after=5,
    )

    _label(doc, "Vision")
    _body(
        doc,
        "Create an offline AI teacher that runs on a regular laptop and helps people learn from trusted content: "
        "Wikipedia, Simple English Wikipedia, dictionaries, textbooks, STEM materials, and curated resources. "
        "Learners ask questions, receive clear explanations, study difficult concepts, understand English more easily, "
        "and continue learning without cloud services or constant connectivity.",
        after=4,
    )

    _label(doc, "Why it matters")
    _body(
        doc,
        "Access to knowledge is one of the strongest tools a person can have. Many people want to learn but lack "
        "stable internet, paid platforms, or personal tutors. USTUD reduces that barrier: a practical companion "
        "for students, self-learners, non-native English speakers, and anyone building their mind with limited resources.",
        after=4,
    )

    _label(doc, "Not a chatbot")
    _body(
        doc,
        "USTUD is an offline learning environment. The AI should not simply answer from memory. It searches local "
        "educational content, retrieves relevant information, explains clearly, and helps the learner understand step "
        "by step. It acts like a patient teacher: simple first, deeper second, encouraging throughout.",
        after=4,
    )

    _label(doc, "First mission")
    _body(
        doc,
        "Build a working MVP. Prove that a user can open USTUD on a Windows laptop, ask a question, search offline "
        "knowledge files, and receive a helpful AI-guided explanation based on retrieved source material. After setup, "
        "the system works without internet access.",
        after=4,
    )

    _label(doc, "Your role")
    _body(
        doc,
        "You have Cursor to build, organize, test, and advance this project. Turn the vision into a working product. "
        "Do not wait for perfect instructions. Create structure, identify blockers, research lightweight tools, and "
        "move forward in clear steps.",
        after=4,
    )

    _label(doc, "First objective — foundation")
    objectives = [
        "Establish the project repository and folder structure.",
        "Research the best lightweight local AI model runner.",
        "Research the best way to search ZIM files (offline Wikipedia).",
        "Create a basic local search proof of concept.",
        "Create a basic local AI response proof of concept.",
        "Connect search and AI into one simple tutoring pipeline.",
        "Build a minimal user interface.",
        "Show the source material used in the answer.",
        "Document every setup step clearly.",
        "Test that the system operates offline after setup.",
    ]
    for item in objectives:
        _bullet(doc, item)

    _body(
        doc,
        "Prioritize practicality over appearance at first. No user accounts, cloud sync, mobile apps, advanced dashboards, "
        "or complex curriculum systems in v1. The first version needs to work.",
        after=5,
    )

    _callout(
        doc,
        "Standard for success",
        "A learner opens USTUD, asks a question, the system searches offline educational material, the AI explains "
        "the answer clearly, and the learner can see where the information came from.",
        "stable",
    )

    _callout(
        doc,
        "Guiding question",
        "Does this help a person learn without depending on the internet? If yes, build it. If not, delay it.",
        "watch",
    )

    _body(
        doc,
        "USTUD is a project about access, education, independence, and human development. Build it with care. "
        "Build it simply. Build it so someone with limited resources can still have a powerful teacher beside them.",
        size=9,
        after=6,
    )

    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("AXIOM  ·  axiom-forhumanity/ustud")
    _run_font(r, FONT_MONO, 7.5, INK_MUTED)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
